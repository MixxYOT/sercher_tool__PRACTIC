"""
Извлечение текста из офисных файлов.
Каждая функция получает путь к файлу, возвращает строку с текстом.
"""
from pathlib import Path
from typing import Tuple

RED = "\033[91m"
RESET = "\033[0m"

# ================= OFFICE =================

def _extract_docx(filepath: str) -> str:
    """
    Извлекает текст из .docx (Microsoft Word 2007+).
    """
    from docx import Document
    
    try:
        doc = Document(filepath)
        paragraphs = []
        
        # Собираем текст из абзацев
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        # Собираем текст из таблиц
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        
        return "\n".join(paragraphs)
    
    except Exception as e:
        return f"[Ошибка чтения docx: {e}]"
    
def _extract_pdf(filepath: str) -> str:
    """
    Извлекает текст из .pdf.
    """
    import fitz

    try: 
        doc = fitz.open(filepath)
        paragraphs = []

        for page in doc:
            text = page.get_text().strip()
            if text:
                paragraphs.append(text)

        doc.close()

        return "\n".join(paragraphs)

    except Exception as e:
        return f"[Ошибка чтения pdf: {e}]"

def _extract_doc(filepath: str) -> str:
    """
    Извлекает текст из .doc (Microsoft Word 97-2003).
    Ищет UTF-16 LE строки напрямую в байтах.
    """
    import olefile
    import re
    
    try:
        ole = olefile.OleFileIO(filepath)
        
        data = b''
        if ole.exists('WordDocument'):
            data = ole.openstream('WordDocument').read()
        
        ole.close()
        
        if not data:
            return "[Поток WordDocument пуст]"
        
        # Ищем ВСЕ читаемые UTF-16 LE строки напрямую в байтах
        # Паттерн: два байта на символ (UTF-16 LE), ищем последовательности букв/цифр/пробелов
        result_parts = []
        
        # Разбиваем данные на пары байт (UTF-16 LE)
        chars = []
        for i in range(0, len(data) - 1, 2):
            # Собираем символ из двух байт (little-endian)
            code = data[i] | (data[i+1] << 8)
            chars.append(code)
        
        # Теперь у нас список Unicode-кодов, фильтруем читаемые
        text_parts = []
        current = []
        
        for code in chars:
            # Проверяем: буква (русская/англ), цифра, пробел, знак препинания
            if (32 <= code <= 126 or           # ASCII печатные
                1040 <= code <= 1103 or        # Русские буквы А-я
                code == 1025 or code == 1105 or # Ё и ё
                code in (10, 13, 9)):          # перенос строки, табуляция
                current.append(chr(code))
            else:
                # Нечитаемый символ — завершаем текущий фрагмент
                if len(current) >= 4:
                    text_parts.append(''.join(current))
                current = []
        
        # Последний фрагмент
        if len(current) >= 4:
            text_parts.append(''.join(current))
        
        # Фильтруем мусор
        clean = []
        for part in text_parts:
            part = part.strip()
            if len(part) < 4:
                continue
            if not re.search(r'[А-Яа-яA-Za-z]', part):
                continue
            # Убираем строки из повторяющихся символов
            if re.match(r'^(.)\1{4,}$', part.replace(' ', '')):
                continue
            clean.append(part)
        
        return '\n'.join(clean) if clean else "[Текст не найден]"
    
    except Exception as e:
        return f"[Ошибка чтения doc: {e}]"
    
def _extract_xlsx(filepath: str) -> str:
    """
    Извлекает текст из .xlsx.
    """
    from openpyxl import load_workbook
    
    try:
        wb = load_workbook(filepath, read_only=True, data_only=True)
        
        all_text = []
        
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            for row in sheet.iter_rows():
                row_text = []
                for cell in row:
                    if cell.value is not None:
                        row_text.append(str(cell.value).strip())
                if row_text:
                    all_text.append(' | '.join(row_text))
        
        wb.close()
        return '\n'.join(all_text)
    
    except Exception as e:
        return f"[Ошибка чтения xlsx: {e}]"

# ================= LIBREOFFICE =================

def _extract_odt(filepath: str) -> str:
    """
    Извлекает текст из .odt (OpenOffice Writer - аналог ворду).
    """
    from odf.opendocument import load
    from odf import text as odf_text
    from odf import teletype
    
    try:
        doc = load(filepath)
        paragraphs = doc.getElementsByType(odf_text.P)
        
        result = []
        for p in paragraphs:
            content = teletype.extractText(p).strip()
            if content:
                result.append(content)
        
        return '\n'.join(result)
    
    except Exception as e:
        return f"[Ошибка чтения odt: {e}]"


def _extract_ods(filepath: str) -> str:
    """
    Извлекает текст из .ods (OpenOffice Calc - аналог exel).
    """
    from odf.opendocument import load
    from odf.table import Table, TableRow, TableCell
    from odf import teletype
    
    try:
        doc = load(filepath)
        tables = doc.getElementsByType(Table)
        
        all_text = []
        
        for table in tables:
            rows = table.getElementsByType(TableRow)
            for row in rows:
                cells = row.getElementsByType(TableCell)
                row_text = []
                for cell in cells:
                    content = teletype.extractText(cell).strip()
                    if content:
                        row_text.append(content)
                if row_text:
                    all_text.append(' | '.join(row_text))
        
        return '\n'.join(all_text)
    
    except Exception as e:
        return f"[Ошибка чтения ods: {e}]"


# ================= ДОП ФУНКЦИИ =================

def _try_read_as_plain_text(filepath: str) -> str:
    """
    Внутренняя функция: пытается прочитать любой файл как простой текст.
    """
    filename = Path(filepath).name
    ext = Path(filepath).suffix.lower()
    
    try:
        with open(filepath, 'r', encoding='utf-8') as file:
            return file.read().strip()
            
    except UnicodeDecodeError:
        try:
            with open(filepath, 'r', encoding='cp1251') as file:
                return file.read().strip()
        except UnicodeDecodeError:
            return False, f"{RED}[Не удалось прочитать '{filename}' как текст. Вероятно, это бинарный формат ({ext}), требующий специального парсера, либо файл поврежден/зашифрован.]{RESET}"
        except Exception as e:
            return False, f"{RED}[Ошибка чтения '{filename}': {e}]{RESET}"
            
    except PermissionError:
        return False, f"{RED}[Нет прав доступа к файлу '{filename}']{RESET}"
    except Exception as e:
        return False, f"{RED}[Ошибка чтения '{filename}': {e}]{RESET}"


# Точка входа для обработки файлов
def extract_text(filepath: str) -> Tuple[bool, str]:
    """
    Определяет тип файла по расширению и вызывает нужный экстрактор.
    """
    ext = Path(filepath).suffix.lower()
    filename = Path(filepath).name
    
    try:
        match ext:
            case '.docx':
                return True, _extract_docx(filepath)
            case '.pdf':
                return True, _extract_pdf(filepath)
            case '.doc':
                return True, _extract_doc(filepath)
            case '.xlsx':
                return True, _extract_xlsx(filepath)
            case '.odt':
                return True, _extract_odt(filepath)
            case '.ods':
                return True, _extract_ods(filepath)
            case _:
                return _try_read_as_plain_text(filepath)
                
    except Exception as e:
        return False, f"{RED}[Критическая ошибка обработки '{filename}': {e}]{RESET}"


# Тесты
if __name__ == "__main__":
    import sys
    
    if len(sys.argv) > 1:
        filepath = sys.argv[1]
    else:
        docs_dir = Path("docs")
        files = list(docs_dir.rglob("*.xlsx"))
        if not files:
            print("Не найдено .xlsx файлов в папке docs/")
            print("Использование: python src/extractors.py путь/к/файлу.xlsx")
            sys.exit(1)
        filepath = str(files[0])
    
    print(f"Файл: {filepath}")
    print(f"Размер: {Path(filepath).stat().st_size} байт\n")
    print("ИЗВЛЕЧЁННЫЙ ТЕКСТ:")
    
    text = extract_text(filepath)
    print(text + '\n')
    
    print(f"Всего символов: {len(text)}")
